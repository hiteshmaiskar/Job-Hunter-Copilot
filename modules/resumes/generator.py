'''
Author:     Sai Vignesh Golla
LinkedIn:   https://www.linkedin.com/in/saivigneshgolla/

Copyright (C) 2024 Sai Vignesh Golla

License:    GNU Affero General Public License
            https://www.gnu.org/licenses/agpl-3.0.en.html
            
GitHub:     https://github.com/GodsScion/Auto_job_applier_linkedIn

Support me: https://github.com/sponsors/GodsScion

'''


import docx
from fpdf import FPDF

def create_resume_docx(user_details, summary, experience, projects, skills, certificates):
    # Create a docx file
    doc = docx.Document()

    # Add user details
    doc.add_heading(user_details['name'], 0)
    doc.add_paragraph(user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'])

    # Add summary
    doc.add_heading('Summary', 1)
    doc.add_paragraph(summary)

    # Add experience
    doc.add_heading('Experience', 1)
    for experience_item in experience:
        doc.add_heading(experience_item['company'], 2)
        doc.add_paragraph(experience_item['role'] + ' | ' + experience_item['dates'])
        doc.add_paragraph(experience_item['achievements'])

    # Add projects
    doc.add_heading('Projects', 1)
    for project in projects:
        doc.add_heading(project['name'], 2)
        doc.add_paragraph(project['description'] + ' | ' + project['technologies'])

    # Add skills
    doc.add_heading('Skills', 1)
    doc.add_paragraph(', '.join(skills))

    # Add certificates
    doc.add_heading('Certificates', 1)
    for certificate in certificates:
        doc.add_heading(certificate['name'], 2)
        doc.add_paragraph(certificate['description'])

    # Save docx file
    doc.save('resume.docx')

    # Create a pdf file
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)

    # Add user details
    pdf.cell(0, 10, user_details['name'], 0, 1, 'C')
    pdf.cell(0, 10, user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'], 0, 1, 'L')

    # Add summary
    pdf.cell(0, 10, 'Summary', 0, 1, 'L')
    pdf.multi_cell(0, 10, summary)

    # Add experience
    pdf.cell(0, 10, 'Experience', 0, 1, 'L')
    for experience_item in experience:
        pdf.cell(0, 10, experience_item['company'], 0, 1, 'L')
        pdf.cell(0, 10, experience_item['role'] + ' | ' + experience_item['dates'], 0, 1, 'L')
        pdf.multi_cell(0, 10, experience_item['achievements'])

    # Add projects
    pdf.cell(0, 10, 'Projects', 0, 1, 'L')
    for project in projects:
        pdf.cell(0, 10, project['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, project['description'] + ' | ' + project['technologies'])

    # Add skills
    pdf.cell(0, 10, 'Skills', 0, 1, 'L')
    pdf.multi_cell(0, 10, ', '.join(skills))

    # Add certificates
    pdf.cell(0, 10, 'Certificates', 0, 1, 'L')
    for certificate in certificates:
        pdf.cell(0, 10, certificate['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, certificate['description'])

    # Save pdf file
    pdf.output('resume.pdf', 'F')


# ══════════════════════════════════════════════════════════════════════════════
# NEW — Resume Tailoring Feature: Tailored PDF Generator using reportlab
# ══════════════════════════════════════════════════════════════════════════════
# The functions below are NEW and independent of the create_resume_docx() /
# FPDF logic above. They use reportlab to generate an ATS-friendly PDF from
# the structured JSON returned by Azure OpenAI.
# ══════════════════════════════════════════════════════════════════════════════

import os

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    _REPORTLAB_AVAILABLE = True
except ImportError:
    _REPORTLAB_AVAILABLE = False


# ──────────────────────────────────────────────────────────────────────────────
# Style helpers
# ──────────────────────────────────────────────────────────────────────────────
def _build_styles():
    """
    Builds and returns a dict of reportlab ParagraphStyle objects
    for a clean, ATS-friendly single-column resume layout.
    """
    base = getSampleStyleSheet()

    styles = {
        # Candidate name — large, bold, centered
        "name": ParagraphStyle(
            name="CandidateName",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        # Contact info line — small, centered, muted
        "contact": ParagraphStyle(
            name="ContactLine",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=6,
            textColor=colors.HexColor("#444444"),
        ),
        # Section heading — bold, dark navy
        "section_heading": ParagraphStyle(
            name="SectionHeading",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            spaceBefore=10,
            spaceAfter=2,
            textColor=colors.HexColor("#1a1a2e"),
        ),
        # Job title / project name — bold
        "job_title": ParagraphStyle(
            name="JobTitle",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=13,
            spaceBefore=6,
            spaceAfter=1,
        ),
        # Dates / location meta line — italic, muted
        "job_meta": ParagraphStyle(
            name="JobMeta",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            leading=12,
            spaceAfter=2,
            textColor=colors.HexColor("#555555"),
        ),
        # Bullet points under experience entries
        "bullet": ParagraphStyle(
            name="BulletPoint",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=13,
            leftIndent=12,
            spaceAfter=1,
        ),
        # General body text — skills, certifications, etc.
        "body": ParagraphStyle(
            name="BodyText",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=13,
            spaceAfter=2,
            alignment=TA_JUSTIFY,
        ),
        # Professional summary — slightly more leading
        "summary": ParagraphStyle(
            name="SummaryText",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=14,
            spaceAfter=4,
            alignment=TA_JUSTIFY,
        ),
    }
    return styles


def _section_heading_flowables(title: str, styles: dict) -> list:
    """Returns a list of flowables: a bold heading + a thin navy rule below it."""
    return [
        Paragraph(title.upper(), styles["section_heading"]),
        HRFlowable(
            width="100%",
            thickness=0.8,
            color=colors.HexColor("#2c3e7a"),
            spaceAfter=4,
        ),
    ]


# ──────────────────────────────────────────────────────────────────────────────
# Main public function
# ──────────────────────────────────────────────────────────────────────────────
def generate_tailored_pdf(resume_data: dict, output_path: str) -> str:
    """
    Renders a tailored resume as a clean, ATS-friendly PDF using reportlab.

    Takes the structured JSON dict returned by azure_tailor_resume() and
    produces a professional single-column PDF resume with sections:
    Name/Contact, Summary, Experience, Skills, Projects, Education, Certifications.

    Args:
        resume_data (dict): Structured resume with keys:
            name, contact (dict), summary, experience (list), skills (list),
            projects (list), education (list), certifications (list).
        output_path (str): Full file path where the PDF should be saved.
            The parent directory is created automatically if it does not exist.
            Example: "all resumes/tailored/3987654321/resume.pdf"

    Returns:
        str: Absolute path to the saved PDF file.

    Raises:
        ImportError: If reportlab is not installed (pip install reportlab).
        ValueError:  If resume_data is None or not a dict.

    Example:
        path = generate_tailored_pdf(
            resume_data,
            "all resumes/tailored/JOB123/resume.pdf"
        )
        print(f"Saved tailored resume to: {path}")
    """
    if not _REPORTLAB_AVAILABLE:
        raise ImportError(
            "reportlab is not installed.\n"
            "Run: pip install reportlab\n"
            "Or inside venv: pip install reportlab"
        )

    if not resume_data or not isinstance(resume_data, dict):
        raise ValueError(
            "resume_data must be a non-empty dict returned by azure_tailor_resume()."
        )

    # Ensure the output directory exists
    abs_output = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(abs_output), exist_ok=True)

    styles = _build_styles()

    # Document with comfortable margins (22mm sides, 18mm top/bottom)
    doc = SimpleDocTemplate(
        abs_output,
        pagesize=A4,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    story = []  # Ordered list of flowable elements

    # ── Header: Candidate Name ────────────────────────────────────────────────
    name = resume_data.get("name", "")
    if name:
        story.append(Paragraph(name, styles["name"]))

    # ── Header: Contact Information ───────────────────────────────────────────
    contact = resume_data.get("contact", {})
    contact_parts = []
    if contact.get("email"):    contact_parts.append(contact["email"])
    if contact.get("phone"):    contact_parts.append(contact["phone"])
    if contact.get("location"): contact_parts.append(contact["location"])
    if contact.get("linkedin"): contact_parts.append(contact["linkedin"])
    if contact.get("github"):   contact_parts.append(contact["github"])
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), styles["contact"]))

    story.append(Spacer(1, 4))

    # ── Section: Professional Summary ────────────────────────────────────────
    summary = resume_data.get("summary", "").strip()
    if summary:
        story += _section_heading_flowables("Professional Summary", styles)
        story.append(Paragraph(summary, styles["summary"]))

    # ── Section: Experience ───────────────────────────────────────────────────
    experience = resume_data.get("experience", [])
    if experience:
        story += _section_heading_flowables("Experience", styles)
        for exp in experience:
            company  = exp.get("company", "")
            title    = exp.get("title", "")
            dates    = exp.get("dates", "")
            location = exp.get("location", "")
            bullets  = exp.get("bullets", [])

            # "Title — Company" on one line
            title_line = (
                f"<b>{title}</b> &mdash; {company}"
                if company else f"<b>{title}</b>"
            )
            story.append(Paragraph(title_line, styles["job_title"]))

            # "Dates | Location" meta line
            meta_parts = []
            if dates:    meta_parts.append(dates)
            if location: meta_parts.append(location)
            if meta_parts:
                story.append(Paragraph(" | ".join(meta_parts), styles["job_meta"]))

            # Bullet points
            for bullet_text in bullets:
                story.append(Paragraph(f"&bull; {bullet_text}", styles["bullet"]))

            story.append(Spacer(1, 3))

    # ── Section: Skills ───────────────────────────────────────────────────────
    skills = resume_data.get("skills", [])
    if skills:
        story += _section_heading_flowables("Skills", styles)
        for skill_line in skills:
            story.append(Paragraph(f"&bull; {skill_line}", styles["body"]))

    # ── Section: Projects ─────────────────────────────────────────────────────
    projects = resume_data.get("projects", [])
    if projects:
        story += _section_heading_flowables("Projects", styles)
        for proj in projects:
            proj_name = proj.get("name", "")
            techs     = proj.get("technologies", "")
            desc      = proj.get("description", "")
            header    = (
                f"<b>{proj_name}</b>" + (f" | <i>{techs}</i>" if techs else "")
            )
            story.append(Paragraph(header, styles["job_title"]))
            if desc:
                story.append(Paragraph(desc, styles["body"]))
            story.append(Spacer(1, 3))

    # ── Section: Education ────────────────────────────────────────────────────
    education = resume_data.get("education", [])
    if education:
        story += _section_heading_flowables("Education", styles)
        for edu in education:
            degree      = edu.get("degree", "")
            institution = edu.get("institution", "")
            dates       = edu.get("dates", "")
            details     = edu.get("details", "")
            title_line  = (
                f"<b>{degree}</b>"
                + (f" &mdash; {institution}" if institution else "")
            )
            story.append(Paragraph(title_line, styles["job_title"]))
            meta_parts = []
            if dates:   meta_parts.append(dates)
            if details: meta_parts.append(details)
            if meta_parts:
                story.append(Paragraph(" | ".join(meta_parts), styles["job_meta"]))
            story.append(Spacer(1, 3))

    # ── Section: Certifications ───────────────────────────────────────────────
    certifications = resume_data.get("certifications", [])
    if certifications:
        story += _section_heading_flowables("Certifications", styles)
        for cert in certifications:
            story.append(Paragraph(f"&bull; {cert}", styles["body"]))

    # ── Build the PDF ─────────────────────────────────────────────────────────
    doc.build(story)
    print(f"[generator] Tailored PDF saved: {abs_output}")
    return abs_output